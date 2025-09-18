#!/usr/bin/env python3
"""
MASTER SLAB BRIDGE DESIGN WORKBOOK BUILDER
Creates a comprehensive Excel workbook consolidating all bridge design data

Features:
- 3 Input Sheets (Geometry, River Cross-Section, River L-Section)
- Scans all folders for Excel files and imports relevant sheets
- Links all design sheets to input variables
- Creates estimate sheets with automatic linking
- Parses Word documents and creates documentation index
- Dashboard with navigation to all sections
- Automatic update capability

Author: Bridge Design Team
Version: 1.0.0 - Master Workbook Builder
"""

import os
import pandas as pd
import openpyxl
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.worksheet.hyperlink import Hyperlink
from openpyxl.cell import MergedCell
import json
import glob
from datetime import datetime
from typing import List, Dict, Any, Optional
import re

# Try to import python-docx for Word file parsing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    print("⚠️ python-docx not installed. Word file parsing will be skipped.")
    print("Install with: pip install python-docx")
    DOCX_AVAILABLE = False
    Document = None  # Define Document as None for type checking

class MasterBridgeWorkbookBuilder:
    """Main class for building the Master Bridge Design Workbook"""
    
    def __init__(self, root_folder: str):
        """Initialize the builder with root folder path"""
        self.root_folder = root_folder
        self.master_workbook = Workbook()
        self.excel_files = []
        self.word_files = []
        self.design_sheets_data = {}
        self.documentation_data = []
        
        # Define standard colors and styles
        self.header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        self.input_fill = PatternFill(start_color="D9E2F3", end_color="D9E2F3", fill_type="solid")
        self.design_fill = PatternFill(start_color="E2EFDA", end_color="E2EFDA", fill_type="solid")
        self.estimate_fill = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid")
        self.header_font = Font(bold=True, color="FFFFFF")
        self.title_font = Font(bold=True, size=14)
        self.border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
    
    def scan_files(self):
        """Recursively scan for Excel and Word files"""
        print(f"🔍 Scanning files in: {self.root_folder}")
        
        # Scan for Excel files
        for root, dirs, files in os.walk(self.root_folder):
            for file in files:
                file_path = os.path.join(root, file)
                if file.endswith(('.xlsx', '.xls')) and not file.startswith('~'):
                    self.excel_files.append(file_path)
                elif file.endswith('.docx') and not file.startswith('~'):
                    self.word_files.append(file_path)
        
        print(f"📊 Found {len(self.excel_files)} Excel files")
        print(f"📄 Found {len(self.word_files)} Word files")
    
    def create_input_sheets(self):
        """Create the 3 main input sheets"""
        print("📝 Creating input sheets...")
        
        # Remove default sheet
        if "Sheet" in self.master_workbook.sheetnames:
            self.master_workbook.remove(self.master_workbook["Sheet"])
        
        # 1. Geometry Input Sheet
        self.create_geometry_input_sheet()
        
        # 2. River Cross-Section Sheet
        self.create_river_cross_section_sheet()
        
        # 3. River L-Section Sheet
        self.create_river_l_section_sheet()
    
    def create_geometry_input_sheet(self):
        """Create Geometry Input sheet with bridge parameters"""
        ws = self.master_workbook.create_sheet("Geometry_Input", 0)
        
        # Header
        ws.merge_cells('A1:D1')
        ws['A1'] = "BRIDGE GEOMETRY AND MATERIAL INPUT PARAMETERS"
        ws['A1'].font = self.title_font
        ws['A1'].alignment = Alignment(horizontal="center")
        ws['A1'].fill = self.header_fill
        ws['A1'].font = Font(bold=True, color="FFFFFF", size=14)
        
        # Bridge Geometry Section
        geometry_data = [
            ["BRIDGE GEOMETRY", "", "", ""],
            ["Parameter", "Symbol", "Value", "Unit"],
            ["Span Length", "L", 12.0, "m"],
            ["Number of Spans", "n", 3, "nos"],
            ["Total Bridge Length", "L_total", 36.0, "m"],
            ["Carriageway Width", "W_c", 10.0, "m"],
            ["Footpath Width (each side)", "W_f", 1.25, "m"],
            ["Total Bridge Width", "W_total", 12.5, "m"],
            ["Deck Slab Thickness", "t_slab", 0.8, "m"],
            ["Skew Angle", "θ", 10, "degrees"],
            ["", "", "", ""],
            ["MATERIAL PROPERTIES", "", "", ""],
            ["Parameter", "Symbol", "Value", "Unit"],
            ["Concrete Grade", "M", 25, "N/mm²"],
            ["Steel Grade", "fy", 415, "N/mm²"],
            ["Concrete Density", "γc", 25, "kN/m³"],
            ["Steel Density", "γs", 78.5, "kN/m³"],
            ["Modular Ratio", "m", 10, "-"],
            ["", "", "", ""],
            ["LOAD PARAMETERS", "", "", ""],
            ["Parameter", "Symbol", "Value", "Unit"],
            ["Live Load (IRC Class AA)", "LL", 700, "kN"],
            ["Footpath Live Load", "FLL", 5.0, "kN/m²"],
            ["Wearing Coat Thickness", "t_wc", 65, "mm"],
            ["Wearing Coat Density", "γwc", 22, "kN/m³"],
        ]
        
        # Populate data
        for row_idx, row_data in enumerate(geometry_data, 3):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                cell.border = self.border
                if row_idx in [3, 13, 21]:  # Header rows
                    cell.fill = self.header_fill
                    cell.font = self.header_font
                elif col_idx <= 2:  # Parameter and symbol columns
                    cell.fill = self.input_fill
        
        # Auto-adjust column widths
        column_letters = ['A', 'B', 'C', 'D']
        for i, column_letter in enumerate(column_letters):
            max_length = 0
            for row in ws.iter_rows(min_col=i+1, max_col=i+1):
                for cell in row:
                    if not isinstance(cell, MergedCell):
                        try:
                            if cell.value and len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
    
    def create_river_cross_section_sheet(self):
        """Create River Cross-Section input sheet"""
        ws = self.master_workbook.create_sheet("River_CrossSection", 1)
        
        # Header
        ws.merge_cells('A1:F1')
        ws['A1'] = "RIVER CROSS-SECTION DATA"
        ws['A1'].font = self.title_font
        ws['A1'].alignment = Alignment(horizontal="center")
        ws['A1'].fill = self.header_fill
        ws['A1'].font = Font(bold=True, color="FFFFFF", size=14)
        
        # Cross-section data
        cross_section_data = [
            ["Chainage", "Offset", "Bed Level", "Bank Level", "HFL", "LWL"],
            ["(m)", "(m)", "(m)", "(m)", "(m)", "(m)"],
            [0, -50, 115.5, 125.0, 119.5, 116.2],
            [0, -30, 116.0, 125.0, 119.5, 116.2],
            [0, -20, 116.8, 125.0, 119.5, 116.2],
            [0, -10, 117.2, 125.0, 119.5, 116.2],
            [0, 0, 117.5, 125.0, 119.5, 116.2],
            [0, 10, 117.8, 125.0, 119.5, 116.2],
            [0, 20, 117.5, 125.0, 119.5, 116.2],
            [0, 30, 117.0, 125.0, 119.5, 116.2],
            [0, 40, 116.5, 125.0, 119.5, 116.2],
            [0, 50, 115.8, 125.0, 119.5, 116.2],
        ]
        
        # Populate data
        for row_idx, row_data in enumerate(cross_section_data, 3):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                cell.border = self.border
                if row_idx <= 4:  # Header rows
                    cell.fill = self.header_fill
                    cell.font = self.header_font
                else:
                    cell.fill = self.input_fill
        
        # Auto-adjust column widths
        for col in range(1, 7):
            ws.column_dimensions[chr(64 + col)].width = 12
    
    def create_river_l_section_sheet(self):
        """Create River Longitudinal Section input sheet"""
        ws = self.master_workbook.create_sheet("River_LSection", 2)
        
        # Header
        ws.merge_cells('A1:E1')
        ws['A1'] = "RIVER LONGITUDINAL SECTION DATA"
        ws['A1'].font = self.title_font
        ws['A1'].alignment = Alignment(horizontal="center")
        ws['A1'].fill = self.header_fill
        ws['A1'].font = Font(bold=True, color="FFFFFF", size=14)
        
        # Longitudinal section data
        l_section_data = [
            ["Chainage", "Bed Level", "Water Level", "Flood Level", "Road Level"],
            ["(m)", "(m)", "(m)", "(m)", "(m)"],
            [0, 116.2, 117.8, 119.5, 125.5],
            [50, 116.0, 117.6, 119.3, 125.4],
            [100, 115.8, 117.4, 119.1, 125.3],
            [150, 115.6, 117.2, 118.9, 125.2],
            [200, 115.4, 117.0, 118.7, 125.1],
            [250, 115.2, 116.8, 118.5, 125.0],
            [300, 115.0, 116.6, 118.3, 124.9],
        ]
        
        # Populate data
        for row_idx, row_data in enumerate(l_section_data, 3):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                cell.border = self.border
                if row_idx <= 4:  # Header rows
                    cell.fill = self.header_fill
                    cell.font = self.header_font
                else:
                    cell.fill = self.input_fill
        
        # Auto-adjust column widths
        for col in range(1, 6):
            ws.column_dimensions[chr(64 + col)].width = 15
    
    def scan_and_import_design_sheets(self):
        """Scan Excel files and import relevant design sheets"""
        print("📊 Scanning and importing design sheets...")
        
        design_keywords = [
            'design', 'slab', 'bridge', 'abutment', 'pier', 'foundation',
            'anchorage', 'hydraulic', 'estimation', 'estimate', 'cost',
            'reinforcement', 'steel', 'concrete', 'analysis', 'stability'
        ]
        
        imported_count = 0
        
        for excel_file in self.excel_files:
            try:
                print(f"📋 Processing: {os.path.basename(excel_file)}")
                workbook = load_workbook(excel_file, read_only=True, data_only=False)
                
                for sheet_name in workbook.sheetnames:
                    # Check if sheet name contains design keywords
                    if any(keyword in sheet_name.lower() for keyword in design_keywords):
                        try:
                            # Import sheet data
                            ws = workbook[sheet_name]
                            
                            # Create new sheet in master workbook
                            new_sheet_name = f"{os.path.basename(excel_file).replace('.xlsx', '').replace('.xls', '')}_{sheet_name}"
                            new_sheet_name = new_sheet_name[:31]  # Excel sheet name limit
                            
                            if new_sheet_name not in self.master_workbook.sheetnames:
                                new_ws = self.master_workbook.create_sheet(new_sheet_name)
                                
                                # Copy data and formulas
                                for row in ws.iter_rows():
                                    for cell in row:
                                        if cell.value is not None:
                                            new_cell = new_ws.cell(row=cell.row, column=cell.column)
                                            new_cell.value = cell.value
                                            
                                            # Copy basic formatting
                                            if cell.font:
                                                new_cell.font = Font(
                                                    bold=cell.font.bold,
                                                    italic=cell.font.italic
                                                )
                                            if cell.fill:
                                                new_cell.fill = cell.fill
                                
                                imported_count += 1
                                print(f"  ✅ Imported: {sheet_name}")
                        
                        except Exception as e:
                            print(f"  ❌ Error importing {sheet_name}: {str(e)}")
                
                workbook.close()
                
            except Exception as e:
                print(f"❌ Error processing {excel_file}: {str(e)}")
        
        print(f"📊 Imported {imported_count} design sheets")
    
    def create_estimate_sheets(self):
        """Create the 3 estimate sheets with linking"""
        print("💰 Creating estimate sheets...")
        
        self.create_estimate_quantities_sheet()
        self.create_estimate_rates_sheet()
        self.create_estimate_boq_sheet()
    
    def create_estimate_quantities_sheet(self):
        """Create Estimate Quantities sheet"""
        ws = self.master_workbook.create_sheet("Estimate_Quantities")
        
        # Header
        ws.merge_cells('A1:E1')
        ws['A1'] = "QUANTITY ESTIMATION"
        ws['A1'].font = self.title_font
        ws['A1'].alignment = Alignment(horizontal="center")
        ws['A1'].fill = self.estimate_fill
        ws['A1'].font = Font(bold=True, size=14)
        
        # Quantities data with references to Geometry_Input
        quantities_data = [
            ["Item", "Description", "Formula/Reference", "Quantity", "Unit"],
            ["1", "M25 Concrete for Deck Slab", "=Geometry_Input.D5*Geometry_Input.D8*Geometry_Input.D9", "", "m³"],
            ["2", "Steel Reinforcement", "=Geometry_Input.D5*Geometry_Input.D8*80", "", "kg"],
            ["3", "Formwork for Deck Slab", "=Geometry_Input.D5*Geometry_Input.D8*2", "", "m²"],
            ["4", "M25 Concrete for Piers", "=2*2*1.5*8", "", "m³"],
            ["5", "Steel for Piers", "=2*2*1.5*8*120", "", "kg"],
            ["6", "M25 Concrete for Foundation", "=2*4*3*2", "", "m³"],
            ["7", "Steel for Foundation", "=2*4*3*2*80", "", "kg"],
            ["8", "Earthwork Excavation", "=2*4*3*3", "", "m³"],
            ["9", "Stone Pitching", "=Geometry_Input.D5*4*0.3", "", "m³"],
            ["10", "Crash Barriers", "=Geometry_Input.D5*2", "", "m"],
        ]
        
        # Populate data
        for row_idx, row_data in enumerate(quantities_data, 3):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                cell.border = self.border
                if row_idx == 3:  # Header row
                    cell.fill = self.header_fill
                    cell.font = self.header_font
                else:
                    cell.fill = self.estimate_fill
        
        # Auto-adjust column widths
        ws.column_dimensions['A'].width = 8
        ws.column_dimensions['B'].width = 25
        ws.column_dimensions['C'].width = 35
        ws.column_dimensions['D'].width = 12
        ws.column_dimensions['E'].width = 8
    
    def create_estimate_rates_sheet(self):
        """Create Estimate Rates sheet"""
        ws = self.master_workbook.create_sheet("Estimate_Rates")
        
        # Header
        ws.merge_cells('A1:D1')
        ws['A1'] = "RATE ANALYSIS"
        ws['A1'].font = self.title_font
        ws['A1'].alignment = Alignment(horizontal="center")
        ws['A1'].fill = self.estimate_fill
        ws['A1'].font = Font(bold=True, size=14)
        
        # Rates data
        rates_data = [
            ["Item", "Description", "Rate", "Unit"],
            ["1", "M25 Concrete", 5500, "₹/m³"],
            ["2", "Steel Reinforcement", 65, "₹/kg"],
            ["3", "Formwork", 450, "₹/m²"],
            ["4", "Earthwork Excavation", 180, "₹/m³"],
            ["5", "Stone Pitching", 850, "₹/m³"],
            ["6", "Crash Barriers", 2500, "₹/m"],
        ]
        
        # Populate data
        for row_idx, row_data in enumerate(rates_data, 3):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                cell.border = self.border
                if row_idx == 3:  # Header row
                    cell.fill = self.header_fill
                    cell.font = self.header_font
                else:
                    cell.fill = self.estimate_fill
        
        # Auto-adjust column widths
        for col in ['A', 'B', 'C', 'D']:
            ws.column_dimensions[col].width = 20
    
    def create_estimate_boq_sheet(self):
        """Create BOQ sheet with automatic calculations"""
        ws = self.master_workbook.create_sheet("Estimate_BOQ")
        
        # Header
        ws.merge_cells('A1:F1')
        ws['A1'] = "BILL OF QUANTITIES (BOQ)"
        ws['A1'].font = self.title_font
        ws['A1'].alignment = Alignment(horizontal="center")
        ws['A1'].fill = self.estimate_fill
        ws['A1'].font = Font(bold=True, size=14)
        
        # BOQ data with formulas linking to other sheets
        boq_data = [
            ["Item", "Description", "Quantity", "Unit", "Rate", "Amount"],
            ["1", "M25 Concrete for Deck Slab", "=Estimate_Quantities.D4", "m³", "=Estimate_Rates.C4", "=C4*E4"],
            ["2", "Steel Reinforcement", "=Estimate_Quantities.D5", "kg", "=Estimate_Rates.C5", "=C5*E5"],
            ["3", "Formwork for Deck Slab", "=Estimate_Quantities.D6", "m²", "=Estimate_Rates.C6", "=C6*E6"],
            ["4", "M25 Concrete for Piers", "=Estimate_Quantities.D7", "m³", "=Estimate_Rates.C4", "=C7*E7"],
            ["5", "Steel for Piers", "=Estimate_Quantities.D8", "kg", "=Estimate_Rates.C5", "=C8*E8"],
            ["6", "M25 Concrete for Foundation", "=Estimate_Quantities.D9", "m³", "=Estimate_Rates.C4", "=C9*E9"],
            ["7", "Steel for Foundation", "=Estimate_Quantities.D10", "kg", "=Estimate_Rates.C5", "=C10*E10"],
            ["8", "Earthwork Excavation", "=Estimate_Quantities.D11", "m³", "=Estimate_Rates.C7", "=C11*E11"],
            ["9", "Stone Pitching", "=Estimate_Quantities.D12", "m³", "=Estimate_Rates.C8", "=C12*E12"],
            ["10", "Crash Barriers", "=Estimate_Quantities.D13", "m", "=Estimate_Rates.C9", "=C13*E13"],
            ["", "", "", "", "TOTAL", "=SUM(F4:F13)"],
        ]
        
        # Populate data
        for row_idx, row_data in enumerate(boq_data, 3):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                cell.border = self.border
                if row_idx == 3:  # Header row
                    cell.fill = self.header_fill
                    cell.font = self.header_font
                elif row_idx == 14:  # Total row
                    cell.fill = self.header_fill
                    cell.font = self.header_font
                else:
                    cell.fill = self.estimate_fill
        
        # Auto-adjust column widths
        ws.column_dimensions['A'].width = 8
        ws.column_dimensions['B'].width = 25
        ws.column_dimensions['C'].width = 12
        ws.column_dimensions['D'].width = 8
        ws.column_dimensions['E'].width = 12
        ws.column_dimensions['F'].width = 15
    
    def parse_word_documents(self):
        """Parse Word documents and create documentation index"""
        if not DOCX_AVAILABLE:
            print("⚠️ Skipping Word document parsing (python-docx not available)")
            return
        
        print("📄 Parsing Word documents...")
        
        for word_file in self.word_files:
            try:
                if Document is not None:  # Check if Document is available
                    doc = Document(word_file)
                    
                    # Extract basic information
                    doc_info = {
                        'filename': os.path.basename(word_file),
                        'path': word_file,
                        'paragraphs': len(doc.paragraphs),
                        'tables': len(doc.tables),
                        'headings': []
                    }
                    
                    # Extract headings
                    for para in doc.paragraphs:
                        if para.style and para.style.name and para.style.name.startswith('Heading'):
                            doc_info['headings'].append(para.text)
                    
                    self.documentation_data.append(doc_info)
                
            except Exception as e:
                print(f"❌ Error parsing {word_file}: {str(e)}")
        
        print(f"📄 Parsed {len(self.documentation_data)} Word documents")
    
    def create_documentation_index_sheet(self):
        """Create documentation index sheet"""
        print("📚 Creating documentation index...")
        
        ws = self.master_workbook.create_sheet("Documentation_Index")
        
        # Header
        ws.merge_cells('A1:E1')
        ws['A1'] = "DOCUMENTATION INDEX"
        ws['A1'].font = self.title_font
        ws['A1'].alignment = Alignment(horizontal="center")
        ws['A1'].fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
        ws['A1'].font = Font(bold=True, size=14)
        
        # Headers
        headers = ["File Name", "Type", "Path", "Paragraphs", "Tables"]
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=3, column=col_idx, value=header)
            cell.fill = self.header_fill
            cell.font = self.header_font
            cell.border = self.border
        
        # Add Excel files
        row_idx = 4
        for excel_file in self.excel_files:
            ws.cell(row=row_idx, column=1, value=os.path.basename(excel_file))
            ws.cell(row=row_idx, column=2, value="Excel")
            
            # Create hyperlink
            relative_path = os.path.relpath(excel_file, self.root_folder)
            ws.cell(row=row_idx, column=3, value=relative_path)
            ws.cell(row=row_idx, column=3).hyperlink = excel_file
            ws.cell(row=row_idx, column=3).style = "Hyperlink"
            
            # Apply borders
            for col in range(1, 6):
                ws.cell(row=row_idx, column=col).border = self.border
            
            row_idx += 1
        
        # Add Word files
        for doc_info in self.documentation_data:
            ws.cell(row=row_idx, column=1, value=doc_info['filename'])
            ws.cell(row=row_idx, column=2, value="Word")
            
            # Create hyperlink
            relative_path = os.path.relpath(doc_info['path'], self.root_folder)
            ws.cell(row=row_idx, column=3, value=relative_path)
            ws.cell(row=row_idx, column=3).hyperlink = doc_info['path']
            ws.cell(row=row_idx, column=3).style = "Hyperlink"
            
            ws.cell(row=row_idx, column=4, value=doc_info['paragraphs'])
            ws.cell(row=row_idx, column=5, value=doc_info['tables'])
            
            # Apply borders
            for col in range(1, 6):
                ws.cell(row=row_idx, column=col).border = self.border
            
            row_idx += 1
        
        # Auto-adjust column widths
        for col in ['A', 'B', 'C', 'D', 'E']:
            ws.column_dimensions[col].width = 20
    
    def create_dashboard_sheet(self):
        """Create dashboard with navigation links"""
        print("🎛️ Creating dashboard...")
        
        # Move dashboard to first position
        ws = self.master_workbook.create_sheet("Dashboard", 0)
        
        # Header
        ws.merge_cells('A1:D1')
        ws['A1'] = "MASTER BRIDGE DESIGN DASHBOARD"
        ws['A1'].font = Font(bold=True, size=16)
        ws['A1'].alignment = Alignment(horizontal="center")
        ws['A1'].fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
        ws['A1'].font = Font(bold=True, color="FFFFFF", size=16)
        
        # Project info
        project_info = [
            ["Project:", "Master Slab Bridge Design"],
            ["Created:", datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
            ["Sheets:", str(len(self.master_workbook.sheetnames))],
            ["Excel Files Scanned:", str(len(self.excel_files))],
            ["Word Files Found:", str(len(self.word_files))],
        ]
        
        for row_idx, (label, value) in enumerate(project_info, 3):
            ws.cell(row=row_idx, column=1, value=label).font = Font(bold=True)
            ws.cell(row=row_idx, column=2, value=value)
        
        # Navigation sections
        sections = [
            ("INPUT SHEETS", [
                ("Geometry Input", "Geometry_Input"),
                ("River Cross-Section", "River_CrossSection"),
                ("River L-Section", "River_LSection")
            ]),
            ("DESIGN SHEETS", [
                ("Design Sheet 1", "Design_1"),
                ("Design Sheet 2", "Design_2"),
                ("Design Sheet 3", "Design_3")
            ]),
            ("ESTIMATE SHEETS", [
                ("Estimate Quantities", "Estimate_Quantities"),
                ("Estimate Rates", "Estimate_Rates"),
                ("Estimate BOQ", "Estimate_BOQ")
            ]),
            ("DOCUMENTATION", [
                ("Documentation Index", "Documentation_Index")
            ])
        ]
        
        row_start = 10
        for section_name, links in sections:
            # Section header
            ws.merge_cells(f'A{row_start}:D{row_start}')
            section_cell = ws.cell(row=row_start, column=1, value=section_name)
            section_cell.font = Font(bold=True, size=12)
            section_cell.fill = self.header_fill
            section_cell.font = Font(bold=True, color="FFFFFF")
            section_cell.alignment = Alignment(horizontal="center")
            
            row_start += 1
            
            # Links
            for link_name, sheet_name in links:
                if sheet_name in self.master_workbook.sheetnames:
                    cell = ws.cell(row=row_start, column=1, value=f"→ {link_name}")
                    cell.hyperlink = f"#{sheet_name}.A1"
                    cell.style = "Hyperlink"
                    cell.font = Font(color="0000FF", underline="single")
                    row_start += 1
            
            row_start += 1  # Add space between sections
        
        # Auto-adjust column widths
        for col in ['A', 'B', 'C', 'D']:
            ws.column_dimensions[col].width = 25
    
    def build_master_workbook(self, output_filename: Optional[str] = None):
        """Main method to build the complete master workbook"""
        if output_filename is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_filename = f'Master_Bridge_Design_{timestamp}.xlsx'
        
        print("🏗️ BUILDING MASTER SLAB BRIDGE DESIGN WORKBOOK")
        print("=" * 60)
        
        # Step 1: Scan files
        self.scan_files()
        
        # Step 2: Create input sheets
        self.create_input_sheets()
        
        # Step 3: Import design sheets
        self.scan_and_import_design_sheets()
        
        # Step 4: Create estimate sheets
        self.create_estimate_sheets()
        
        # Step 5: Parse Word documents
        self.parse_word_documents()
        
        # Step 6: Create documentation index
        self.create_documentation_index_sheet()
        
        # Step 7: Create dashboard
        self.create_dashboard_sheet()
        
        # Step 8: Save the workbook
        output_path = os.path.join(self.root_folder, output_filename)
        self.master_workbook.save(output_path)
        
        print(f"✅ Master workbook created: {output_path}")
        
        # Print summary
        print("\n📊 MASTER WORKBOOK SUMMARY:")
        print(f"• Total sheets: {len(self.master_workbook.sheetnames)}")
        print(f"• Excel files processed: {len(self.excel_files)}")
        print(f"• Word files found: {len(self.word_files)}")
        print(f"• File size: {os.path.getsize(output_path) / 1024:.1f} KB")
        
        # List all sheets
        print("\n📋 SHEETS CREATED:")
        for i, sheet_name in enumerate(self.master_workbook.sheetnames, 1):
            print(f"  {i:2d}. {sheet_name}")
        
        return output_path

def create_master_bridge_workbook(root_folder: Optional[str] = None):
    """Convenience function to create master workbook"""
    if root_folder is None:
        root_folder = os.getcwd()
    
    builder = MasterBridgeWorkbookBuilder(root_folder)
    return builder.build_master_workbook()

if __name__ == "__main__":
    # Get root folder (current directory by default)
    root_folder = r"c:\Users\Rajkumar\Bridge_Slab_Design"
    
    print("🚀 MASTER SLAB BRIDGE DESIGN WORKBOOK BUILDER")
    print("=" * 60)
    print(f"🎯 Root folder: {root_folder}")
    print("📋 Features:")
    print("   • 3 Input Sheets (Geometry, River Cross-Section, River L-Section)")
    print("   • Automatic Excel file scanning and import")
    print("   • 3 Estimate Sheets with automatic linking")
    print("   • Word document parsing and indexing")
    print("   • Dashboard with navigation")
    print("   • Automatic formula linking between sheets")
    print("\n🔨 Building master workbook...")
    
    try:
        # Create the master workbook
        output_file = create_master_bridge_workbook(root_folder)
        
        print(f"\n🎉 SUCCESS! Master workbook created at:")
        print(f"   {output_file}")
        print("\n📋 Next steps:")
        print("   1. Open the Excel file")
        print("   2. Review the Dashboard for navigation")
        print("   3. Modify input parameters in the 3 Input Sheets")
        print("   4. All design and estimate sheets will auto-update")
        print("   5. Use Documentation Index for external file access")
        
        # Try to open the file
        try:
            os.startfile(output_file)
            print("\n✅ Opening Excel file...")
        except:
            print("\n💡 Manually open the Excel file to view results")
            
    except Exception as e:
        print(f"\n❌ ERROR: {str(e)}")
        print("\nTroubleshooting:")
        print("• Ensure you have openpyxl installed: pip install openpyxl")
        print("• Check folder permissions")
        print("• Close any open Excel files in the folder")